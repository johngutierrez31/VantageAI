#!/usr/bin/env python3
"""
Multi-Format Policy Document Generator - Static File Processing Tool

Generates professional policy documents in DOCX, HTML, and PDF formats from
customized policy JSON.

IMPORTANT - STATIC ANALYSIS ONLY:
This script processes static policy data files. It does NOT:
- Connect to live systems or networks
- Execute code or make network requests
- Access databases or servers
- Require credentials or system access

Analyzes: Customized policy JSON files
Purpose: Generate formatted policy documents in multiple output formats

Requirements:
- python-docx (for DOCX generation)
- markdown2 (for HTML generation)
- weasyprint (for PDF generation - optional)
"""

import json
import sys
import os
from typing import Dict, Any, List
from datetime import datetime


def check_dependencies():
    """Check for required Python packages."""
    missing = []

    try:
        import docx
    except ImportError:
        missing.append("python-docx")

    try:
        import markdown2
    except ImportError:
        missing.append("markdown2")

    # WeasyPrint is optional
    try:
        import weasyprint
    except ImportError:
        pass  # PDF generation will be disabled but not required

    if missing:
        print("Error: Missing required dependencies")
        print(f"Please install: pip install {' '.join(missing)}")
        sys.exit(1)


def load_json(filepath: str) -> Dict[str, Any]:
    """Load JSON file."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        print(f"Error: File '{filepath}' not found")
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"Error: Invalid JSON in '{filepath}': {e}")
        sys.exit(1)


def generate_docx(policy: Dict[str, Any], output_path: str):
    """Generate DOCX document."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Set document properties
    customization = policy.get('customization', {})
    core_properties = doc.core_properties
    core_properties.title = policy.get('title', 'Policy Document')
    core_properties.author = customization.get('organization', 'Unknown')
    core_properties.subject = 'Cybersecurity Policy'

    # Title
    title = doc.add_heading(policy.get('title', 'Untitled Policy'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata table
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'

    metadata_fields = [
        ('Company', customization.get('organization', 'N/A')),
        ('Version', customization.get('version', '1.0')),
        ('Effective Date', customization.get('effective_date', 'TBD')),
        ('Review Schedule', customization.get('review_schedule', 'Annually')),
        ('Responsible Officer', customization.get('responsible_officer', 'N/A')),
        ('Contact', f"{customization.get('organization', 'N/A')} Security Team")
    ]

    for i, (label, value) in enumerate(metadata_fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        # Bold the label
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()  # Spacing

    # Description
    description = policy.get('description', '')
    if description:
        p = doc.add_paragraph(description)
        p.paragraph_format.space_after = Pt(12)

    # Sections
    sections = policy.get('sections', [])
    for section in sections:
        _add_section_to_docx(doc, section, level=1)

    # Compliance frameworks
    frameworks = customization.get('frameworks', [])
    regulations = customization.get('regulations', [])

    if frameworks or regulations:
        doc.add_page_break()
        doc.add_heading('Compliance Frameworks', 1)

        if frameworks:
            doc.add_heading('Frameworks', 2)
            for framework in frameworks:
                doc.add_paragraph(framework, style='List Bullet')

        if regulations and regulations != ["None"]:
            doc.add_heading('Regulatory Requirements', 2)
            for regulation in regulations:
                if regulation != "None":
                    doc.add_paragraph(regulation, style='List Bullet')

    # Approval footer
    doc.add_page_break()
    doc.add_heading('Approval and Review', 1)

    approval_table = doc.add_table(rows=3, cols=2)
    approval_table.style = 'Light Grid Accent 1'

    approval_data = [
        ('Approved by', customization.get('responsible_officer', 'N/A')),
        ('Approval Date', customization.get('effective_date', 'TBD')),
        ('Next Review Date', _calculate_next_review(customization))
    ]

    for i, (label, value) in enumerate(approval_data):
        row = approval_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].bold = True

    # Footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"{customization.get('organization', 'Organization')} - {policy.get('title', 'Policy')} - Version {customization.get('version', '1.0')}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    doc.save(output_path)


def _add_section_to_docx(doc, section: Dict[str, Any], level: int = 1):
    """Add section to DOCX document."""
    from docx.shared import Pt

    # Section title
    title = section.get('title', '')
    if title:
        doc.add_heading(title, level)

    # Section content
    content = section.get('content', '')
    if content:
        # Split by paragraphs
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                # Check if it's a list item
                if para_text.strip().startswith('- ') or para_text.strip().startswith('* '):
                    # List item
                    lines = para_text.split('\n')
                    for line in lines:
                        if line.strip().startswith('- ') or line.strip().startswith('* '):
                            item_text = line.strip()[2:]
                            doc.add_paragraph(item_text, style='List Bullet')
                        else:
                            doc.add_paragraph(line.strip())
                else:
                    p = doc.add_paragraph(para_text.strip())
                    p.paragraph_format.space_after = Pt(6)

    # Subsections
    subsections = section.get('subsections', [])
    for subsection in subsections:
        _add_section_to_docx(doc, subsection, level + 1)


def generate_html(policy: Dict[str, Any], output_path: str):
    """Generate HTML document."""
    import markdown2

    customization = policy.get('customization', {})

    # Build Markdown content (reuse Markdown generator logic)
    markdown_content = _policy_to_markdown(policy)

    # Convert to HTML
    html_body = markdown2.markdown(
        markdown_content,
        extras=['tables', 'fenced-code-blocks', 'header-ids', 'toc']
    )

    # Wrap in HTML template
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{policy.get('title', 'Policy Document')}</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #34495e;
            border-bottom: 2px solid #95a5a6;
            padding-bottom: 5px;
            margin-top: 30px;
        }}
        h3 {{
            color: #555;
            margin-top: 20px;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }}
        th, td {{
            padding: 12px;
            text-align: left;
            border: 1px solid #ddd;
        }}
        th {{
            background-color: #3498db;
            color: white;
        }}
        tr:nth-child(even) {{
            background-color: #f2f2f2;
        }}
        .metadata {{
            background-color: #ecf0f1;
            padding: 15px;
            border-radius: 5px;
            margin-bottom: 30px;
        }}
        .metadata strong {{
            color: #2c3e50;
        }}
        hr {{
            border: none;
            border-top: 2px solid #bdc3c7;
            margin: 30px 0;
        }}
        ul {{
            margin-left: 20px;
        }}
        li {{
            margin-bottom: 8px;
        }}
        .footer {{
            margin-top: 50px;
            padding-top: 20px;
            border-top: 2px solid #95a5a6;
            font-size: 0.9em;
            color: #7f8c8d;
            text-align: center;
        }}
        @media print {{
            body {{
                max-width: 100%;
            }}
        }}
    </style>
</head>
<body>
    {html_body}
    <div class="footer">
        <p>Generated by Cybersecurity Policy Generator</p>
        <p>&copy; {datetime.now().year} {customization.get('organization', 'Organization')}. All rights reserved.</p>
    </div>
</body>
</html>
"""

    # Save
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html)


def generate_pdf(policy: Dict[str, Any], output_path: str, html_path: str = None):
    """Generate PDF document from HTML."""
    try:
        from weasyprint import HTML
    except ImportError:
        print("Warning: weasyprint not installed. PDF generation disabled.")
        print("Install with: pip install weasyprint")
        return False

    # If HTML path provided, use it; otherwise generate temporary HTML
    if html_path and os.path.exists(html_path):
        HTML(filename=html_path).write_pdf(output_path)
    else:
        # Generate temporary HTML
        import tempfile
        with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as tmp:
            tmp_path = tmp.name
            generate_html(policy, tmp_path)

        # Convert to PDF
        HTML(filename=tmp_path).write_pdf(output_path)

        # Clean up
        os.unlink(tmp_path)

    return True


def _policy_to_markdown(policy: Dict[str, Any]) -> str:
    """Convert policy to Markdown (simplified version)."""
    customization = policy.get('customization', {})
    md = f"""# {policy.get('title', 'Untitled Policy')}

**Company:** {customization.get('organization', 'N/A')}
**Version:** {customization.get('version', '1.0')}
**Effective Date:** {customization.get('effective_date', 'TBD')}
**Review Schedule:** {customization.get('review_schedule', 'Annually')}
**Responsible Officer:** {customization.get('responsible_officer', 'N/A')}

---

"""

    # Description
    description = policy.get('description', '')
    if description:
        md += f"{description}\n\n"

    # Sections
    sections = policy.get('sections', [])
    for section in sections:
        md += _section_to_markdown(section)

    # Compliance
    frameworks = customization.get('frameworks', [])
    if frameworks:
        md += "\n## Compliance Frameworks\n\n"
        for framework in frameworks:
            md += f"- {framework}\n"

    return md


def _section_to_markdown(section: Dict[str, Any], level: int = 2) -> str:
    """Convert section to Markdown."""
    md = ""

    title = section.get('title', '')
    if title:
        md += f"{'#' * level} {title}\n\n"

    content = section.get('content', '')
    if content:
        md += f"{content}\n\n"

    subsections = section.get('subsections', [])
    for subsection in subsections:
        md += _section_to_markdown(subsection, level + 1)

    return md


def _calculate_next_review(customization: Dict[str, Any]) -> str:
    """Calculate next review date."""
    effective_date = customization.get('effective_date', '')
    review_schedule = customization.get('review_schedule', 'Annually')

    if not effective_date:
        return "TBD"

    try:
        from dateutil.relativedelta import relativedelta
        eff_date = datetime.fromisoformat(effective_date)

        if review_schedule == "Quarterly":
            next_review = eff_date + relativedelta(months=3)
        elif review_schedule == "Semi-annually":
            next_review = eff_date + relativedelta(months=6)
        elif review_schedule == "Annually":
            next_review = eff_date + relativedelta(years=1)
        elif review_schedule == "Bi-annually":
            next_review = eff_date + relativedelta(years=2)
        else:
            return "As needed"

        return next_review.strftime('%Y-%m-%d')

    except (ImportError, ValueError):
        return f"{review_schedule} from {effective_date}"


def main():
    """Main entry point."""
    import argparse

    parser = argparse.ArgumentParser(
        description='Generate policy documents in DOCX, HTML, and PDF formats',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Generate all formats
  python generate_docx_html_pdf.py \\
    --input customized_policy.json \\
    --output-dir ./output \\
    --formats docx html pdf

  # Generate only DOCX
  python generate_docx_html_pdf.py \\
    --input customized_policy.json \\
    --output-dir ./output \\
    --formats docx

  # Custom output filenames
  python generate_docx_html_pdf.py \\
    --input customized_policy.json \\
    --output-dir ./policies \\
    --basename "AcceptableUsePolicy" \\
    --formats docx html

Requirements:
  pip install python-docx markdown2 weasyprint
        """
    )

    parser.add_argument(
        '--input',
        required=True,
        help='Path to customized policy JSON file'
    )

    parser.add_argument(
        '--output-dir',
        default='./output',
        help='Output directory for generated files (default: ./output)'
    )

    parser.add_argument(
        '--basename',
        help='Base filename for outputs (default: derived from policy title)'
    )

    parser.add_argument(
        '--formats',
        nargs='+',
        choices=['docx', 'html', 'pdf'],
        default=['docx', 'html', 'pdf'],
        help='Formats to generate (default: all)'
    )

    parser.add_argument(
        '--verbose',
        action='store_true',
        help='Show detailed processing information'
    )

    args = parser.parse_args()

    # Check dependencies
    check_dependencies()

    # Load policy
    if args.verbose:
        print(f"Loading policy from: {args.input}")

    policy = load_json(args.input)

    # Determine base filename
    if args.basename:
        basename = args.basename
    else:
        # Derive from policy title
        title = policy.get('title', 'Policy')
        basename = title.replace(' ', '')  # Remove spaces

    # Create output directory
    os.makedirs(args.output_dir, exist_ok=True)

    print(f"\n{'='*80}")
    print(f"Generating Policy Documents")
    print(f"{'='*80}\n")
    print(f"Policy: {policy.get('title', 'Unknown')}")
    print(f"Organization: {policy.get('customization', {}).get('organization', 'N/A')}")
    print(f"Output Directory: {os.path.abspath(args.output_dir)}\n")

    generated_files = []

    # Generate DOCX
    if 'docx' in args.formats:
        docx_path = os.path.join(args.output_dir, f"{basename}.docx")
        if args.verbose:
            print(f"Generating DOCX: {docx_path}")
        generate_docx(policy, docx_path)
        generated_files.append(('DOCX', docx_path))
        print(f"✓ DOCX generated: {basename}.docx")

    # Generate HTML
    html_path = None
    if 'html' in args.formats:
        html_path = os.path.join(args.output_dir, f"{basename}.html")
        if args.verbose:
            print(f"Generating HTML: {html_path}")
        generate_html(policy, html_path)
        generated_files.append(('HTML', html_path))
        print(f"✓ HTML generated: {basename}.html")

    # Generate PDF
    if 'pdf' in args.formats:
        pdf_path = os.path.join(args.output_dir, f"{basename}.pdf")
        if args.verbose:
            print(f"Generating PDF: {pdf_path}")

        success = generate_pdf(policy, pdf_path, html_path)
        if success:
            generated_files.append(('PDF', pdf_path))
            print(f"✓ PDF generated: {basename}.pdf")
        else:
            print(f"⚠ PDF generation skipped (weasyprint not installed)")

    # Summary
    print(f"\n{'='*80}")
    print(f"Generation Complete")
    print(f"{'='*80}\n")
    print(f"Generated {len(generated_files)} file(s):\n")

    for format_type, filepath in generated_files:
        file_size = os.path.getsize(filepath)
        print(f"  [{format_type}] {os.path.basename(filepath)} ({file_size:,} bytes)")

    print(f"\n{'='*80}\n")


if __name__ == "__main__":
    main()
